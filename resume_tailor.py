"""
Resume tailoring via DeepSeek and .docx upload to Google Drive.
Fails gracefully — pipeline continues if tailoring is unavailable.
"""

import os
import re
from datetime import datetime
from pathlib import Path

from dotenv import load_dotenv

load_dotenv()

TAILOR_ENABLED = False

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload
    from openai import OpenAI

    from google_drive_auth import get_drive_credentials

    TAILOR_ENABLED = True
except ImportError as exc:
    print(f"Warning: resume tailoring dependencies missing ({exc}) — tailoring disabled.")

OUTPUTS_DIR = Path("outputs")
SECTION_HEADERS = {"PROFILE", "SKILLS", "EXPERIENCE", "EDUCATION"}
PLACEHOLDER_PATTERNS = [
    r"\[YOUR NAME\]",
    r"\[Phone\]",
    r"\[Email\]",
    r"\[LinkedIn\]",
    r"\[University Name\]",
    r"\[Add your",
    r"Bullet point \d",
]


def has_placeholders(text: str) -> bool:
    """Return True if text still contains template placeholders."""
    return any(re.search(pattern, text, re.IGNORECASE) for pattern in PLACEHOLDER_PATTERNS)


def _parse_domain_mismatch(content: str) -> tuple[bool, str]:
    """Return (is_mismatch, reason) if the LLM flagged an incompatible domain."""
    if not re.search(r"DOMAIN_MISMATCH_FLAG:\s*true", content, re.IGNORECASE):
        return False, ""

    reason_match = re.search(r"REASON:\s*(.+)", content, re.IGNORECASE | re.DOTALL)
    if reason_match:
        reason = reason_match.group(1).strip().splitlines()[0].strip()
        return True, reason or "Domain mismatch — no credible bridge to required industry."
    return True, "Domain mismatch — no credible bridge to required industry."


def _clean_llm_resume(text: str) -> str:
    """Remove AI preamble and ATS notes — keep only the resume body."""
    if "--- ATS NOTES" in text:
        text = text.split("--- ATS NOTES", 1)[0]

    lines = text.splitlines()
    cleaned = []
    started = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if started:
                cleaned.append("")
            continue

        if not started:
            lower = stripped.lower()
            if lower.startswith(("here is ", "below is ", "output ", "tailored resume")):
                continue
            if (
                stripped.upper() in SECTION_HEADERS
                or re.match(r"^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+", stripped)
                or "@" in stripped
                or "+44" in stripped
                or re.match(r"^\+?\d", stripped)
            ):
                started = True

        if started:
            cleaned.append(stripped)

    return "\n".join(cleaned).strip()


def _strip_html(text: str) -> str:
    if not text:
        return ""
    clean = re.sub(r"<[^>]+>", " ", text)
    return re.sub(r"\s+", " ", clean).strip()


def tailor_resume(base_resume_text: str, job: dict) -> str | None:
    """Call DeepSeek to tailor resume for one job. Returns plain text or None."""
    if not TAILOR_ENABLED:
        return None

    api_key = os.getenv("DEEPSEEK_API_KEY")
    if not api_key:
        print("Resume tailoring: DEEPSEEK_API_KEY not set.")
        return None

    title = job.get("title", "Unknown")

    prompt = f"""
You are a senior hiring manager and CV strategist with 15 years of
experience shortlisting candidates for analyst, CRM, and business
intelligence roles at UK companies. You have reviewed over 10,000 CVs.

You know exactly why good candidates get rejected:
- Safe, forgettable language that sounds like everyone else
- Activity-based bullets with no outcomes
- Skills sections that are keyword dumps
- Thin experience padded with weak bullets
- No differentiation in the profile

Your job is to rewrite this candidate's resume for the specific job
below. The output must make a hiring manager pause and read properly —
not skim and move on.

═══════════════════════════════════════════════
DOMAIN COMPATIBILITY CHECK — perform this before writing anything:
═══════════════════════════════════════════════

Compare the job's required domain (industry, function, technical
discipline) against the candidate's actual work history domains.

If the job is in a SPECIALISED REGULATED INDUSTRY the candidate has
zero direct or adjacent experience in — examples: investment banking,
trading operations, payments/fintech infrastructure, healthcare
clinical roles, legal practice, pharmaceutical regulatory affairs,
aerospace/defence engineering — and the candidate's background has no
genuine bridge to it (not even via transferable skills with a credible
story), do NOT write a tailored resume claiming relevance.

Instead, output exactly this in place of the resume:

DOMAIN_MISMATCH_FLAG: true
REASON: [One sentence explaining the specific gap — e.g. 'This role
requires direct sales and trading or payments industry experience;
candidate's background is in graduate engagement and data analysis
with no financial services exposure, and the stated salary band
confirms this is a specialist-level hire, not an entry point.']

Do not attempt to write a tailored resume for jobs flagged this way.
A resume cannot manufacture domain credibility that doesn't exist, and
attempting to do so produces a CV that reads as inauthentic to any
experienced hiring manager in that field, which damages the candidate's
credibility for every future application from the same CV style if
they reuse phrasing patterns.

If the domain gap is merely adjacent/transferable (not zero bridge),
proceed with the rewrite using honest Problem 11 framing — only use
DOMAIN_MISMATCH_FLAG for genuinely incompatible specialised industries.

═══════════════════════════════════════════════
THE 11 PROBLEMS YOU MUST FIX IN THIS RESUME:
═══════════════════════════════════════════════

PROBLEM 1 — DUPLICATE EMPLOYER, SAME DATES
If the resume shows two roles at the same employer with identical or
overlapping dates, combine them into one entry. Use format:
"[Title 1] / [Title 2], [Company] | [Date]"
with a single set of bullets that covers both responsibilities.
This removes the impression of experience padding.

PROBLEM 2 — SKILLS SECTION IS A DUMP
Do not list more than 12-14 skills. Group them into 2-3 natural
categories without using category headers (just ordering and line
breaks signal the grouping):
- Line 1: core technical tools (BI, data, analysis tools)
- Line 2: domain and process skills
- Line 3: platforms and productivity (only if genuinely strong)
Front-load the skills most relevant to this specific job.
Remove any skill that cannot be demonstrated by at least one
bullet in the experience section.

PROBLEM 3 — WEAK ROLE GETTING TOO MANY BULLETS
Any role that is clearly non-analytical (retail, admin, basic intern)
gets a maximum of 2 bullets. These 2 bullets must find the most
business-relevant angle possible — inventory = stock data management,
customer service = stakeholder communication, training = knowledge
transfer. If there is no credible business angle, cut to 1 bullet.

PROBLEM 4 — ACTIVITY BULLETS WITH NO OUTCOMES
Every bullet in the strongest 1-2 roles must answer: "so what?"
If the original bullet describes what was done but not what changed,
add the implied outcome. Use this test:
Original: "Built Power BI dashboards to monitor survey performance"
Ask: did anything improve? did anyone use it to make decisions?
was there a frequency (daily, weekly)? was there a scale (N users)?
If yes to any of these, add it. If genuinely unknown, use language
that implies impact without fabricating: "enabling the team to..."
"informing weekly decisions on..." "used by senior managers to..."

PROBLEM 5 — BULLETS TOO UNIFORM IN LENGTH
Break the pattern. Aim for this distribution per role:
- 1 short punchy bullet (under 12 words)
- 2 medium bullets (15-20 words)
- 1 detailed bullet with context and outcome (25-30 words)
Never have two consecutive bullets of the same length or structure.

PROBLEM 6 — PROFILE TOO GENERIC
The profile must contain:
- One specific, concrete number or achievement from the experience
- The exact job title being applied for (from the JD)
- One domain-specific phrase that signals genuine expertise
- A forward-looking line that connects the candidate's background
  to what the employer is trying to achieve
Length: 3 sentences maximum. No "I". No "passionate about".
No "looking to leverage my skills". Direct, factual, confident.

PROBLEM 7 — OLD SHORT INTERNSHIP OVER-BULLLETED
Any internship under 3 months or over 3 years ago gets 2 bullets
maximum. Make them tight. Cut anything that does not add new
evidence of a skill.

PROBLEM 8 — MBA UNDERUSED
In the Education section, add one line under the MBA entry:
Relevant modules or focus areas that overlap with this job.
Keep it to one line. Only add if there is a genuine connection
(e.g. for BA roles: "Modules in business strategy, data-driven
decision making, and organisational analytics").
If the candidate has a dissertation or project relevant to the role,
mention it in one clause.

PROBLEM 9 — NO PROJECTS/PORTFOLIO MENTIONED
If the base resume has no projects section, add a short one after
Education with this instruction: add a placeholder that says:
[Add 1-2 personal projects here — e.g. a Power BI dashboard, a
Python automation tool, or a data analysis competition]
Only do this if the section would not look empty — if there are
genuinely no projects to add, skip this section entirely and do
not create an empty one.

PROBLEM 10 — "PRESENT" ON A NON-PRIMARY ROLE
If a role shows "Present" as the end date and it appears to be
a part-time or secondary role (retail, bar work, etc.), change
the framing to make clear it is part-time/ongoing alongside job
search. Add "(part-time)" in brackets after the title. This
prevents the hiring manager from assuming this is the main job.

PROBLEM 11 — JD PHRASE MIRRORING
If the original candidate experience is meaningfully different in
DOMAIN from what the JD describes (e.g. candidate did engagement
tracking, JD wants system migration validation), do not force
identical language. Use honest framing instead: name what the
candidate actually did, and let the transferable skill speak for
itself, rather than relabelling the activity to sound like the JD's
exact responsibility. A reconciliation check on engagement records is
not the same discipline as source-to-target migration validation —
describe it honestly as the former, and let the hiring manager draw
the (legitimate) connection rather than asserting equivalence.

═══════════════════════════════════════════════
ATS KEYWORD RULES (apply alongside the above):
═══════════════════════════════════════════════
- Extract the top 8 keywords from the job description
- Every keyword must appear at least once in the resume
- At least 3 must appear in bullets (not just the skills section)
- Use the exact phrase the JD uses, not a synonym
  (ATS does exact and near-exact matching, not semantic)
- Justification test: only add a keyword if at least one bullet
  in the original resume demonstrates the underlying skill
- Skills injected must be defensible in an interview

═══════════════════════════════════════════════
WRITING RULES — MUST FOLLOW ALL OF THESE:
═══════════════════════════════════════════════
- British English throughout (analyse not analyze, etc.)
- No AI phrases: never use spearheaded, leveraged, orchestrated,
  facilitated, championed, fostered, utilised, synergies, robust,
  cutting-edge, dynamic, results-driven, detail-oriented, passionate,
  proven track record, seamlessly, streamlined, pivotal, innovative
- Active verbs that are specific: built, ran, analysed, produced,
  managed, wrote, tracked, reported, presented, trained, cut, grew
- Vary sentence openings — no two consecutive bullets starting with
  the same word
- Read every bullet aloud — if it sounds like a corporate press
  release, rewrite it in plain English first, then add the keyword
- Specific beats generic always: "1,600 graduates" not "large cohort",
  "daily" not "regular", "senior managers" not "key stakeholders"
- Numbers that already exist in the resume: never change them
- Job titles at employers: never change them
- Dates: never change them
- Do not invent percentages or figures that are not in the original

CRITICAL — AVOID JD MIRRORING:
Do not copy 3+ consecutive words directly from the job description
into the resume, even if rephrased elsewhere. Do not reuse a specific
number, percentage, or target figure from the JD unless that exact
figure already existed in the candidate's original resume before this
rewrite. If the JD states a target like '99.9% accuracy' and the
candidate's original resume had no such figure, do NOT invent one to
match it — this is fabrication, not tailoring, and experienced hiring
managers recognise this pattern immediately as it is the single
biggest tell of an AI-generated, reverse-engineered CV.

Test before finalising: read the job description and the rewritten
resume side by side. If 3 or more bullets independently echo the same
sequence of words or the same structure as the JD's responsibility
list (e.g. JD says 'design and execute data validation checks' and the
resume says 'designed and executed data validation checks'), rewrite
those bullets using different sentence structure and your own framing
of what the candidate actually did. Mirroring JD structure across
multiple bullets is a pattern hiring managers are now trained to spot
— it reads as templated, not as genuine experience.

The goal is alignment in SUBSTANCE (the underlying skill is genuinely
transferable) without mirroring in SURFACE FORM (identical phrasing,
identical structure, identical figures).

═══════════════════════════════════════════════
OUTPUT FORMAT — EXACT STRUCTURE REQUIRED:
═══════════════════════════════════════════════

[Candidate Full Name]
[Phone] | [Email] | [LinkedIn] | [Location]

PROFILE
[3 sentences max. No "I". Specific. One number. One JD keyword.]

SKILLS
[Line 1: core technical — max 6-7, most JD-relevant first]
[Line 2: domain/process — max 4-5]
[Line 3: platforms — max 3-4, only if strong]

EXPERIENCE

[If two roles same employer same dates: combine as Title1 / Title2]
[Company | Date Range]
- [short punchy bullet — under 12 words]
- [medium bullet with keyword — 15-20 words]
- [medium bullet with outcome — 15-20 words]
- [detailed bullet with context + outcome — 25-30 words]

[Next role]
[Company | Date Range]
- [max 2 bullets if non-analytical role]

EDUCATION
[Degree — Institution (Year)]
[One line: relevant modules or focus if applicable]
[Second degree — Institution (Year)]

--- ATS NOTES (REMOVE BEFORE SENDING) ---
Job targeted: [exact title]
Keywords injected: [keyword → which bullet]
Problems fixed: [list which of the 11 problems were addressed]
Bullets rewritten: [original opening → new opening → reason]
Skills added: [skill → justified by]
Interview prep flag: [any line that needs the candidate to be ready
to explain — e.g. "if asked about Power Query, be ready to describe
the specific dashboards built at UEL"]
DOMAIN GAP CHECK: [Honest one-line assessment of whether the
candidate's actual experience is in the same discipline as this job,
or merely adjacent/transferable. Example: 'Candidate has engagement
data reconciliation experience, not system migration validation —
this is an adjacent skill, be ready to address this gap directly in
any screening call.']

═══════════════════════════════════════════════
INPUTS:
═══════════════════════════════════════════════
JOB TITLE: {job.get('title', '')}
COMPANY: {job.get('company', {}).get('display_name', '') if isinstance(job.get('company'), dict) else job.get('company', '')}

JOB DESCRIPTION:
{str(job.get('description', ''))[:2500]}

CANDIDATE BASE RESUME:
{base_resume_text[:3500]}

NOW OUTPUT THE FULL REWRITTEN RESUME. FOLLOW EVERY RULE ABOVE.
DO NOT SKIP ANY OF THE 11 PROBLEMS. DO NOT USE AI PHRASES.
"""

    try:
        client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=2500,
        )
        content = response.choices[0].message.content.strip()
        mismatch, reason = _parse_domain_mismatch(content)
        if mismatch:
            company = (
                job.get("company", {}).get("display_name", "Unknown")
                if isinstance(job.get("company"), dict)
                else job.get("company", "Unknown")
            )
            print(f"Skipped tailoring (domain mismatch): {title} at {company} — {reason}")
            job["_domain_mismatch"] = True
            job["_domain_mismatch_reason"] = reason
            return None
        cleaned = _clean_llm_resume(content)
        if has_placeholders(cleaned):
            print(f"Resume tailoring rejected placeholder output for '{title}'.")
            return None
        if "EXPERIENCE" not in cleaned.upper():
            print(f"Resume tailoring rejected incomplete output for '{title}'.")
            return None
        return cleaned
    except Exception as exc:
        print(f"Resume tailoring error for '{title}': {exc}")
        return None


def write_cover_letter(base_resume_text: str, job: dict) -> str | None:
    """Generate a tailored UK cover letter as plain text."""
    if not TAILOR_ENABLED:
        return None

    api_key = os.getenv("DEEPSEEK_API_KEY")
    if not api_key:
        return None

    title = job.get("title", "Unknown")
    company = (
        job.get("company", {}).get("display_name", "")
        if isinstance(job.get("company"), dict)
        else job.get("company", "")
    )
    description = _strip_html(job.get("description", ""))[:2000]
    today = datetime.now().strftime("%d %B %Y")

    prompt = f"""Write a one-page UK cover letter for this job application.

RULES:
- British English throughout
- Use ONLY facts from the candidate resume — do not invent experience
- 3–4 short paragraphs, direct tone, no buzzwords
- No placeholders like [Your Name] — use real details from the resume
- Address: Dear Hiring Manager
- Sign off: Yours sincerely, then the candidate's real name
- Mention the specific role title and company name
- Explain 2–3 relevant achievements from the resume that match this job
- Keep total length under 320 words

OUTPUT FORMAT (plain text only, no markdown):
[Candidate Full Name]
[Phone] | [Email] | [LinkedIn] | [Location]

{today}

Hiring Manager
{company}

Dear Hiring Manager,

[paragraphs]

Yours sincerely,
[Candidate Full Name]

JOB TITLE: {title}
COMPANY: {company}
JOB DESCRIPTION:
{description}

CANDIDATE RESUME:
{base_resume_text[:3000]}
"""

    try:
        client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
            max_tokens=1200,
        )
        content = response.choices[0].message.content.strip()
        if has_placeholders(content):
            print(f"Cover letter rejected placeholder output for '{title}'.")
            return None
        if "Dear Hiring Manager" not in content:
            print(f"Cover letter rejected incomplete output for '{title}'.")
            return None
        return content
    except Exception as exc:
        print(f"Cover letter error for '{title}': {exc}")
        return None


def _job_filename_parts(job: dict) -> tuple[str, str, str]:
    company = (
        str(job.get("company", {}).get("display_name", ""))
        if isinstance(job.get("company"), dict)
        else str(job.get("company", "Company"))
    )
    title = str(job.get("title", "Role"))
    date_str = datetime.now().strftime("%Y%m%d")
    safe_company = re.sub(r"[^a-zA-Z0-9]", "_", company)[:20]
    safe_title = re.sub(r"[^a-zA-Z0-9]", "_", title)[:25]
    return safe_company, safe_title, date_str


def _build_cover_letter_docx(letter_text: str) -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    for line in letter_text.splitlines():
        stripped = _strip_markdown(line)
        if not stripped:
            continue
        para = doc.add_paragraph(stripped)
        for run in para.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)

    return doc


def _upload_docx_to_drive(local_path: Path, filename: str) -> str | None:
    folder_id = os.getenv("GOOGLE_DOCS_FOLDER_ID", "")
    if not folder_id:
        print("Google Drive upload: missing GOOGLE_DOCS_FOLDER_ID.")
        return None

    creds = get_drive_credentials()
    if not creds:
        return None

    try:
        service = build("drive", "v3", credentials=creds)
        metadata = {"name": filename, "parents": [folder_id]}
        media = MediaFileUpload(
            str(local_path),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        uploaded = (
            service.files()
            .create(
                body=metadata,
                media_body=media,
                fields="id, webViewLink",
                supportsAllDrives=True,
            )
            .execute()
        )
        service.permissions().create(
            fileId=uploaded["id"],
            body={"type": "anyone", "role": "reader"},
        ).execute()
        return uploaded.get("webViewLink")
    except Exception as exc:
        print(f"Google Drive upload error for '{filename}': {exc}")
        return None


def save_and_upload_cover_letter(letter_text: str, job: dict) -> str | None:
    """Build cover letter .docx, save locally, upload to Drive."""
    if not TAILOR_ENABLED:
        return None

    safe_company, safe_title, date_str = _job_filename_parts(job)
    filename = f"cover_letter_{safe_company}_{safe_title}_{date_str}.docx"

    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    local_path = OUTPUTS_DIR / filename

    try:
        doc = _build_cover_letter_docx(letter_text)
        doc.save(local_path)
        job["_cover_letter_local_path"] = str(local_path)
    except Exception as exc:
        print(f"Failed to build cover letter for '{job.get('title')}': {exc}")
        return None

    return _upload_docx_to_drive(local_path, filename)


DOC_SECTION_HEADERS = {"PROFILE", "SKILLS", "EXPERIENCE", "EDUCATION", "PROJECTS"}
DATE_PATTERN = re.compile(r"20\d{2}")
ARTEFACT_LINES = {"•", "*", "**"}


def _strip_markdown(line: str) -> str:
    text = line.strip()
    text = re.sub(r"^#+\s*", "", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", text)
    text = text.replace("**", "").replace("*", "")
    return text.strip()


def apply_bottom_border(para):
    p_pr = para._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def apply_top_border(para):
    p_pr = para._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "000000")
    p_bdr.append(top)
    p_pr.append(p_bdr)


def add_paragraph(
    doc,
    text="",
    *,
    runs=None,
    font_name="Calibri",
    font_size=10.5,
    bold=False,
    italic=False,
    color=None,
    alignment=None,
    space_before=None,
    space_after=None,
    left_indent=None,
    first_line_indent=None,
    line_spacing=1.0,
):
    para = doc.add_paragraph()
    if alignment is not None:
        para.alignment = alignment

    pf = para.paragraph_format
    pf.line_spacing = line_spacing
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if left_indent is not None:
        pf.left_indent = left_indent
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent

    if runs:
        for run_spec in runs:
            run = para.add_run(run_spec.get("text", ""))
            run.font.name = run_spec.get("font_name", font_name)
            run.font.size = Pt(run_spec.get("font_size", font_size))
            run.font.bold = run_spec.get("bold", bold)
            run.font.italic = run_spec.get("italic", italic)
            if run_spec.get("color") is not None:
                run.font.color.rgb = run_spec["color"]
    else:
        run = para.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        if color is not None:
            run.font.color.rgb = color

    return para


def _is_contact_line(cleaned: str) -> bool:
    return "@" in cleaned or "+44" in cleaned or "linkedin.com" in cleaned.lower()


def _is_section_header(cleaned: str, original: str) -> bool:
    upper = cleaned.upper().strip("- ").strip()
    if upper in DOC_SECTION_HEADERS:
        return True
    if "ATS NOTES" in upper:
        return True
    return original.strip().startswith("---") or cleaned.startswith("---")


def _is_job_title_line(cleaned: str, original: str) -> bool:
    if cleaned.startswith("•") or cleaned.startswith("-"):
        return False
    if " | " not in cleaned:
        return False
    if DATE_PATTERN.search(cleaned):
        return True
    return original.strip().endswith("**")


def _is_bullet_line(cleaned: str, original: str) -> bool:
    original_stripped = original.strip()
    cleaned_stripped = cleaned.strip()
    return (
        original_stripped.startswith("-")
        or original_stripped.startswith("•")
        or cleaned_stripped.startswith("-")
        or cleaned_stripped.startswith("•")
    )


def _is_skills_line(cleaned: str, current_section: str | None) -> bool:
    return current_section == "SKILLS" and " | " in cleaned and not DATE_PATTERN.search(cleaned)


def _next_meaningful_line(line_pairs: list[tuple[str, str]], start_idx: int) -> tuple[int, str] | None:
    for idx in range(start_idx, len(line_pairs)):
        _, cleaned = line_pairs[idx]
        if cleaned and cleaned not in ARTEFACT_LINES:
            return idx, cleaned
    return None


def _default_candidate_name() -> str:
    resume_path = Path("config/resume.txt")
    if resume_path.exists():
        for line in resume_path.read_text(encoding="utf-8").splitlines():
            cleaned = _strip_markdown(line)
            if cleaned and not _is_contact_line(cleaned) and cleaned.upper() not in DOC_SECTION_HEADERS:
                return cleaned
    return "Candidate"


def _resolve_header_lines(line_pairs: list[tuple[str, str]]) -> tuple[str, str | None, set[int]]:
    """Find candidate name and contact from lines before the first section header."""
    name = None
    contact = None
    skip_indices: set[int] = set()

    for idx, (original, cleaned) in enumerate(line_pairs):
        if not cleaned or cleaned in ARTEFACT_LINES:
            continue
        if _is_section_header(cleaned, original):
            break
        if _is_contact_line(cleaned):
            if contact is None:
                contact = cleaned
                skip_indices.add(idx)
        elif name is None:
            name = cleaned
            skip_indices.add(idx)

    if not name or _is_contact_line(name):
        name = _default_candidate_name()

    return name, contact, skip_indices


def _build_docx(tailored_text: str) -> Document:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    line_pairs = [(original, _strip_markdown(original)) for original in tailored_text.splitlines()]

    candidate_name, contact_line, header_skip = _resolve_header_lines(line_pairs)

    add_paragraph(
        doc,
        candidate_name,
        font_size=16,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(4),
    )
    if contact_line:
        add_paragraph(
            doc,
            contact_line,
            font_size=10,
            bold=False,
            color=RGBColor(85, 85, 85),
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=Pt(8),
        )

    current_section = None
    ats_notes_mode = False
    first_ats_content = True
    previous_was_bullet = False

    idx = 0
    while idx < len(line_pairs):
        original, cleaned = line_pairs[idx]

        if idx in header_skip:
            idx += 1
            continue

        if cleaned in ARTEFACT_LINES:
            idx += 1
            continue

        if "[Add 1-2 personal projects" in cleaned:
            add_paragraph(
                doc,
                cleaned,
                font_size=9,
                italic=True,
                color=RGBColor(160, 160, 160),
                space_after=Pt(1),
            )
            idx += 1
            continue

        if not cleaned:
            next_line = _next_meaningful_line(line_pairs, idx + 1)
            if previous_was_bullet and next_line:
                _, next_cleaned = next_line
                if _is_job_title_line(next_cleaned, line_pairs[next_line[0]][0]):
                    add_paragraph(doc, "", space_after=Pt(3))
            previous_was_bullet = False
            idx += 1
            continue

        if _is_section_header(cleaned, original):
            if "ATS NOTES" in cleaned.upper() or original.strip().startswith("---"):
                ats_notes_mode = True
                first_ats_content = True

            header_text = cleaned.strip("- ").strip()
            para = add_paragraph(
                doc,
                header_text,
                font_size=11,
                bold=True,
                space_before=Pt(10),
                space_after=Pt(2),
            )
            apply_bottom_border(para)

            section_key = cleaned.upper().strip("- ").split("(")[0].strip()
            if section_key in DOC_SECTION_HEADERS:
                current_section = section_key
            previous_was_bullet = False
            idx += 1
            continue

        if ats_notes_mode:
            para = add_paragraph(
                doc,
                cleaned,
                font_size=9,
                italic=True,
                color=RGBColor(136, 136, 136),
                space_after=Pt(1),
            )
            if first_ats_content:
                apply_top_border(para)
                first_ats_content = False
            previous_was_bullet = False
            idx += 1
            continue

        if _is_skills_line(cleaned, current_section):
            add_paragraph(
                doc,
                cleaned,
                font_size=10.5,
                bold=False,
                space_after=Pt(3),
            )
            previous_was_bullet = False
            idx += 1
            continue

        if _is_job_title_line(cleaned, original):
            role_part, date_part = cleaned.split(" | ", 1)
            add_paragraph(
                doc,
                runs=[
                    {
                        "text": role_part,
                        "font_size": 10.5,
                        "bold": True,
                    },
                    {
                        "text": f" | {date_part}",
                        "font_size": 10,
                        "bold": False,
                        "color": RGBColor(85, 85, 85),
                    },
                ],
                space_before=Pt(6),
                space_after=Pt(1),
            )
            previous_was_bullet = False
            idx += 1
            continue

        if _is_bullet_line(cleaned, original):
            bullet_text = cleaned.lstrip("•-").strip()
            add_paragraph(
                doc,
                f"• {bullet_text}",
                font_size=10.5,
                bold=False,
                left_indent=Cm(0.4),
                first_line_indent=Cm(-0.4),
                space_after=Pt(2),
            )
            previous_was_bullet = True
            idx += 1
            continue

        add_paragraph(
            doc,
            cleaned,
            font_size=10.5,
            bold=False,
            space_after=Pt(3),
        )
        previous_was_bullet = False
        idx += 1

    return doc


def save_and_upload_resume(tailored_text: str, job: dict) -> str | None:
    """Build .docx, save locally, upload to Drive. Returns webViewLink or None."""
    if not TAILOR_ENABLED:
        return None

    safe_company, safe_title, date_str = _job_filename_parts(job)
    filename = f"resume_{safe_company}_{safe_title}_{date_str}.docx"

    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    local_path = OUTPUTS_DIR / filename

    try:
        doc = _build_docx(tailored_text)
        doc.save(local_path)
        job["_cv_local_path"] = str(local_path)
    except Exception as exc:
        print(f"Failed to build .docx for '{job.get('title')}': {exc}")
        return None

    return _upload_docx_to_drive(local_path, filename)
